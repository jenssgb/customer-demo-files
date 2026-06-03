from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT

ROOT = Path(__file__).resolve().parent
DATA = ROOT / "data"
DATA.mkdir(parents=True, exist_ok=True)

FONT = "Aptos"


def set_font(document: Document) -> None:
    styles = document.styles
    styles["Normal"].font.name = FONT
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = FONT
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 2"].font.name = FONT
    styles["Heading 2"].font.size = Pt(13)


def add_cover(document: Document, title: str, subtitle: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    document.add_paragraph(subtitle).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("Demo document - fictitious data for Microsoft 365 Copilot legal workflow demonstration.").alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()


def add_clause(document: Document, number: str, title: str, body: str) -> None:
    document.add_heading(f"{number}. {title}", level=2)
    document.add_paragraph(body)


def create_contract() -> None:
    document = Document()
    set_font(document)
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    add_cover(document, "Master Services Agreement", "Between Contoso Real Estate Holding GmbH and Northwind Property Services Ltd.")
    document.add_heading("Agreement Summary", level=1)
    document.add_paragraph("This Master Services Agreement governs property management, facilities coordination, tenant service operations, reporting, and related advisory services for a mixed real estate portfolio in Germany, France, the Netherlands, the United Kingdom, and Singapore.")
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Value"
    rows = [
        ("Effective date", "1 July 2026"),
        ("Initial term", "36 months, automatic renewal for additional 24 months unless terminated 180 days before expiry"),
        ("Annual fee", "EUR 4,850,000 plus pass-through costs"),
        ("Governing law", "England and Wales"),
        ("Portfolio", "72 commercial and mixed-use properties"),
        ("Service credits", "Maximum 3 percent of monthly service fees"),
    ]
    for field, value in rows:
        cells = table.add_row().cells
        cells[0].text = field
        cells[1].text = value

    document.add_heading("Commercial and Legal Terms", level=1)
    add_clause(document, "1", "Scope of Services", "Northwind shall provide property operations support, vendor coordination, facilities ticket triage, tenant communication, service charge reporting, lease administration support, ESG data collection, and monthly portfolio performance reporting. Northwind may subcontract any part of the services without prior written approval where it believes such subcontracting is operationally efficient.")
    add_clause(document, "2", "Service Levels", "Northwind shall use commercially reasonable efforts to meet the service levels in Schedule 2. Failure to meet a service level for three consecutive months may result in a service credit, provided Contoso gives written notice within five business days after the relevant reporting period. Total service credits are capped at three percent of the monthly service fee and are Contoso's sole and exclusive remedy for service level failures.")
    add_clause(document, "3", "Fees and Indexation", "Fees are payable monthly in advance. Northwind may increase fees once per calendar year by the higher of CPI, wage inflation in the service location, or five percent. Pass-through costs require no prior approval if they relate to urgent operational needs or health and safety matters.")
    add_clause(document, "4", "Data Protection", "Each party shall comply with applicable data protection law. Northwind may process tenant, visitor, building access, service ticket, vendor, and energy usage data for service delivery, analytics, benchmarking, product improvement, and other legitimate business purposes. Northwind may transfer personal data to affiliates and subcontractors located outside the European Economic Area where appropriate safeguards are in place.")
    add_clause(document, "5", "Confidentiality", "Each party shall protect confidential information using reasonable care. Confidentiality obligations expire three years after disclosure. Northwind may use aggregated and anonymized portfolio data to create industry benchmarks and commercial insights, provided no individual tenant name is disclosed.")
    add_clause(document, "6", "Audit Rights", "Contoso may audit Northwind once per year upon 30 days' notice during normal business hours. Audits must not disrupt Northwind operations and must be limited to records directly related to services performed for Contoso. Northwind may charge reasonable support fees for audit assistance.")
    add_clause(document, "7", "Liability", "Northwind's total aggregate liability under this Agreement shall not exceed the fees paid in the three months preceding the event giving rise to the claim. Neither party shall be liable for indirect, consequential, special, punitive, or loss of profit damages, even if advised of the possibility of such damages. The liability cap applies to all claims including confidentiality, data protection, security incidents, and indemnities.")
    add_clause(document, "8", "Indemnity", "Northwind shall indemnify Contoso against third-party claims arising from Northwind's gross negligence or willful misconduct. Contoso shall indemnify Northwind for tenant claims, property defects, legacy environmental conditions, and instructions given by Contoso personnel.")
    add_clause(document, "9", "Termination", "Either party may terminate for material breach if the breach remains uncured for 60 days after written notice. Contoso may terminate for convenience after the first 18 months by paying an early termination fee equal to nine months of average fees plus committed pass-through costs. Termination assistance is available for an additional fee.")
    add_clause(document, "10", "Change Control", "Any change in portfolio size, reporting requirements, service level scope, ESG data obligations, regulatory requirements, or system integrations shall be treated as a chargeable change request unless expressly included in Schedule 1.")
    add_clause(document, "11", "Dispute Resolution", "Senior executives shall first attempt to resolve disputes informally. If unresolved after 45 days, disputes shall be referred to binding arbitration in London under LCIA rules. Proceedings shall be confidential and conducted in English.")
    add_clause(document, "12", "AI and Automation", "Northwind may use AI tools, workflow automation, and subcontractor platforms to provide the services. Northwind remains responsible for outputs it approves for delivery to Contoso. Contoso acknowledges that AI-generated drafts, summaries, and recommendations may contain inaccuracies and must be reviewed by qualified personnel before reliance.")

    document.add_heading("Schedule 2 - Selected Service Levels", level=1)
    slas = document.add_table(rows=1, cols=4)
    slas.style = "Table Grid"
    slas.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = slas.rows[0].cells
    hdr[0].text = "Service"
    hdr[1].text = "Target"
    hdr[2].text = "Measurement"
    hdr[3].text = "Credit"
    for row in [
        ("Critical tenant issue triage", "4 business hours", "Ticket timestamp to first response", "0.5% monthly fee"),
        ("Monthly portfolio report", "Business day 8", "Report sent to Contoso", "0.25% monthly fee"),
        ("ESG meter-data completeness", "92%", "Meter points with validated reading", "0.25% monthly fee"),
        ("Vendor invoice coding", "95% within 5 business days", "Invoice workflow data", "0.5% monthly fee"),
    ]:
        cells = slas.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    document.save(DATA / "Northwind_Property_Services_Agreement.docx")


def create_playbook() -> None:
    document = Document()
    set_font(document)
    add_cover(document, "Contoso Legal Playbook", "Service Agreements and Property Operations Contracts")
    document.add_heading("Review Standard", level=1)
    document.add_paragraph("This playbook defines Contoso's preferred fallback positions for service agreements involving property operations, facilities management, tenant service support, data processing, reporting, and ESG data collection.")
    document.add_heading("Clause Positions", level=1)
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Topic", "Preferred position", "Fallback", "Escalate if"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    rows = [
        ("Subcontracting", "Prior written consent for material subcontractors; flow-down obligations required", "Notice plus right to object for critical services", "Unrestricted subcontracting or no flow-down"),
        ("Service credits", "Credits are not sole remedy; chronic failure permits termination", "Sole remedy only for minor SLA misses", "Credits capped below 5% or notice window under 10 business days"),
        ("Indexation", "Annual increase capped at 3% or CPI, whichever is lower", "CPI plus documented wage inflation, capped at 5%", "Provider can choose highest of multiple indices without cap"),
        ("Data protection", "Processing limited to service delivery and documented instructions", "Analytics allowed only with anonymization and written approval", "Product improvement or international transfer is broad or undefined"),
        ("Confidentiality", "At least 5 years; trade secrets indefinite", "3 years only for non-sensitive commercial information", "All confidentiality expires after 3 years"),
        ("Audit", "Annual audit at no charge; additional audit after incident", "Reasonable support fees only for extraordinary requests", "Provider controls scope too narrowly or charges routine audit support"),
        ("Liability", "12 months fees cap; uncapped data protection, confidentiality, fraud, willful misconduct", "6 months cap; carve-outs preserved", "3 months cap or cap covers data/security/confidentiality"),
        ("Termination", "30-day cure for material breach; convenience termination without punitive fee after transition", "60-day cure only for non-critical breaches", "Early termination fee exceeds 3 months fees"),
        ("Governing law", "Local law of main service location or mutually neutral EU forum", "England and Wales acceptable only with EU data protection safeguards", "Forum creates material enforcement or cost disadvantage"),
        ("AI use", "AI use disclosed; no confidential data in external AI without approval; human review mandatory", "Approved AI tools only; audit trail retained", "Broad AI use with no tool, data, or review controls"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    document.add_heading("Negotiation Instructions", level=1)
    instructions = [
        "Preserve tracked changes when proposing edits.",
        "Add comments where business owner input is needed.",
        "Do not provide legal advice as final determination; create review notes for counsel.",
        "Prioritize data protection, liability, audit, and termination provisions.",
        "For AI use clauses, require approved tools, human review, and no broad product-improvement use of confidential data.",
    ]
    for instruction in instructions:
        document.add_paragraph(instruction, style="List Bullet")
    document.save(DATA / "Contoso_Legal_Playbook_Service_Agreements.docx")


def create_issue_memo() -> None:
    document = Document()
    set_font(document)
    add_cover(document, "Counterparty Position Memo", "Northwind negotiation context for the property services agreement")
    document.add_heading("Commercial Context", level=1)
    document.add_paragraph("Northwind is the incumbent property operations provider for 29 of the 72 properties. The business wants a fast signature before the new reporting cycle starts on 1 July 2026. Procurement is comfortable with the annual fee but has concerns about indexation, pass-through costs, and termination fees.")
    document.add_heading("Known Business Concerns", level=1)
    for item in [
        "Facilities team wants flexibility to add ESG reporting without a full change request each time.",
        "Finance wants stronger audit support during month-end service charge reconciliation.",
        "IT security wants approval rights over AI tooling and subcontractor platforms.",
        "Asset management can accept England and Wales law if data protection and audit controls are strengthened.",
        "CFO will not accept a three-month liability cap if it also covers confidentiality or data incidents.",
    ]:
        document.add_paragraph(item, style="List Bullet")
    document.add_heading("Requested Output", level=1)
    document.add_paragraph("Legal should produce a clause risk summary, negotiation redlines, and a concise executive recommendation for the signature meeting.")
    document.save(DATA / "Counterparty_Position_Memo.docx")


if __name__ == "__main__":
    create_contract()
    create_playbook()
    create_issue_memo()
    print("Generated legal demo Word assets:")
    for path in sorted(DATA.glob("*.docx")):
        print(f" - {path.name}")
